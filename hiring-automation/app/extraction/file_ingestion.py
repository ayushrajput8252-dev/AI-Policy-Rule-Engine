"""
File ingestion: turns an uploaded PDF/DOCX into raw text plus structural
signals (tables/images present) and extraction-quality warnings.

Primary extractors:
  - PDF  -> pdfplumber (page.extract_text + page.find_tables + page.images)
  - DOCX -> python-docx (paragraphs + tables + inline shapes)

Fallback:
  - PDF  -> PyMuPDF (fitz), used only when pdfplumber yields suspiciously
    little text (e.g. pdfplumber occasionally under-extracts certain PDFs;
    PyMuPDF has a different text layout engine and often recovers more).

We never silently accept a bad extraction: `likely_bad_extraction` and a
list of `warnings` are always returned alongside the text so callers (and
the ATS formatting-penalty step) can react instead of scoring garbage.
"""
import io
import re
from dataclasses import dataclass, field
from typing import List

import pdfplumber
from docx import Document as DocxDocument

try:
    import fitz  # PyMuPDF
    _HAS_PYMUPDF = True
except ImportError:  # pragma: no cover - optional dependency
    _HAS_PYMUPDF = False

MIN_CHARS_OK = 200          # below this, extraction is considered thin
MIN_CHARS_HARD_FAIL = 20    # below this, extraction is considered failed
MIN_ALPHA_RATIO = 0.4       # garbled-text heuristic (too few letters vs. total chars)


@dataclass
class ExtractionResult:
    text: str
    file_type: str                 # "pdf" | "docx"
    extraction_method: str         # "pdfplumber" | "pdfplumber+pymupdf_fallback" | "python-docx"
    has_tables: bool = False
    has_images: bool = False
    likely_bad_extraction: bool = False
    warnings: List[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.text)


def _alpha_ratio(text: str) -> float:
    stripped = re.sub(r"\s+", "", text)
    if not stripped:
        return 0.0
    alpha = sum(1 for c in stripped if c.isalpha())
    return alpha / len(stripped)


def _flag_quality(text: str, warnings: List[str]) -> bool:
    """Appends warnings in place; returns True if extraction looks bad."""
    bad = False
    if len(text.strip()) < MIN_CHARS_HARD_FAIL:
        warnings.append(
            "extraction_failed: little to no text extracted — file is likely a "
            "scanned image, empty, or corrupted"
        )
        bad = True
    elif len(text.strip()) < MIN_CHARS_OK:
        warnings.append(
            "extraction_thin: unusually little text extracted — result may be incomplete"
        )
        bad = True

    if text.strip() and _alpha_ratio(text) < MIN_ALPHA_RATIO:
        warnings.append(
            "extraction_garbled: extracted text has an unusually low letter ratio — "
            "possible encoding issue or scanned/image-based content"
        )
        bad = True

    return bad


def extract_from_pdf(file_bytes: bytes) -> ExtractionResult:
    warnings: List[str] = []
    text_parts = []
    has_tables = False
    has_images = False

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
            if not has_tables:
                try:
                    has_tables = len(page.find_tables()) > 0
                except Exception:
                    pass
            if not has_images:
                has_images = len(page.images) > 0

    text = "\n".join(text_parts).strip()
    method = "pdfplumber"

    # Fallback to PyMuPDF if pdfplumber came back thin/empty and PyMuPDF is available.
    if len(text) < MIN_CHARS_OK and _HAS_PYMUPDF:
        try:
            fallback_text = _extract_pdf_with_pymupdf(file_bytes)
            if len(fallback_text.strip()) > len(text):
                text = fallback_text.strip()
                method = "pdfplumber+pymupdf_fallback"
                warnings.append(
                    "pdfplumber_thin_result: pdfplumber extracted little text; "
                    "used PyMuPDF fallback which recovered more content"
                )
        except Exception as exc:  # pragma: no cover - defensive
            warnings.append(f"pymupdf_fallback_failed: {exc}")

    bad = _flag_quality(text, warnings)
    if not text.strip():
        warnings.append(
            "possible_scanned_pdf: no extractable text layer found — this file may "
            "need OCR before it can be parsed"
        )

    return ExtractionResult(
        text=text,
        file_type="pdf",
        extraction_method=method,
        has_tables=has_tables,
        has_images=has_images,
        likely_bad_extraction=bad,
        warnings=warnings,
    )


def _extract_pdf_with_pymupdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def extract_from_docx(file_bytes: bytes) -> ExtractionResult:
    warnings: List[str] = []
    doc = DocxDocument(io.BytesIO(file_bytes))

    paragraphs = [p.text for p in doc.paragraphs]
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            table_texts.append(" | ".join(cell.text for cell in row.cells))

    text = "\n".join(paragraphs)
    if table_texts:
        text += "\n" + "\n".join(table_texts)
    text = text.strip()

    has_tables = len(doc.tables) > 0
    has_images = _docx_has_images(doc)

    bad = _flag_quality(text, warnings)

    return ExtractionResult(
        text=text,
        file_type="docx",
        extraction_method="python-docx",
        has_tables=has_tables,
        has_images=has_images,
        likely_bad_extraction=bad,
        warnings=warnings,
    )


def _docx_has_images(doc: DocxDocument) -> bool:
    try:
        return len(doc.inline_shapes) > 0
    except Exception:  # pragma: no cover - defensive
        return False


def extract_text(filename: str, file_bytes: bytes) -> ExtractionResult:
    """Dispatch on file extension. Raises ValueError for unsupported types."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_from_pdf(file_bytes)
    if lower.endswith(".docx"):
        return extract_from_docx(file_bytes)
    raise ValueError(
        f"Unsupported file type for '{filename}'. Only .pdf and .docx are supported."
    )
