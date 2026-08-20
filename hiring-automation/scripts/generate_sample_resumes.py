"""
Generates sample resume files (one DOCX, one PDF) into sample_resumes/ so
the API can be smoke-tested without needing real resumes on hand.

Run: python scripts/generate_sample_resumes.py
"""
import os

import fitz  # PyMuPDF
from docx import Document

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "sample_resumes")

RESUME_TEXT_LINES = [
    "John Doe",
    "john.doe@example.com | (415) 555-2671 | linkedin.com/in/johndoe | github.com/johndoe",
    "",
    "Summary",
    "Backend engineer with 6+ years of experience building scalable APIs.",
    "",
    "Experience",
    "Senior Backend Engineer, Acme Corp",
    "Jan 2021 - Present",
    "- Built microservices in Python (FastAPI, Django) deployed on AWS with Docker and Kubernetes.",
    "- Designed PostgreSQL schemas and Redis caching layers handling 10k req/s.",
    "- Set up CI/CD pipelines using GitHub Actions and Terraform for infrastructure as code.",
    "",
    "Backend Engineer, Beta Inc",
    "Jun 2018 - Dec 2020",
    "- Developed REST APIs with Node.js and Express.js, integrated MongoDB and Kafka.",
    "- Wrote unit tests with Pytest and Jest, improved coverage from 40% to 85%.",
    "",
    "Education",
    "Bachelor of Science in Computer Science, University of Washington, 2018",
    "",
    "Skills",
    "Python, JavaScript, TypeScript, FastAPI, Django, Node.js, React, PostgreSQL, MongoDB,",
    "Redis, Docker, Kubernetes, AWS, Terraform, GitHub Actions, Machine Learning, Git",
]


def generate_docx():
    doc = Document()
    for line in RESUME_TEXT_LINES:
        doc.add_paragraph(line)

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Certification"
    table.rows[0].cells[1].text = "AWS Certified Solutions Architect"
    table.rows[1].cells[0].text = "Certification"
    table.rows[1].cells[1].text = "Certified Kubernetes Administrator"

    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, "john_doe.docx")
    doc.save(path)
    print(f"Wrote {path}")


def generate_pdf():
    doc = fitz.open()
    page = doc.new_page()
    text = "\n".join(RESUME_TEXT_LINES)
    page.insert_textbox(fitz.Rect(36, 36, 560, 780), text, fontsize=10)

    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, "john_doe.pdf")
    doc.save(path)
    doc.close()
    print(f"Wrote {path}")


def generate_thin_pdf():
    """A near-empty PDF to exercise the bad-extraction detection path."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_textbox(fitz.Rect(36, 36, 560, 100), "Resume", fontsize=10)
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, "thin_scanned_like.pdf")
    doc.save(path)
    doc.close()
    print(f"Wrote {path}")


# --- Additional fixtures for /score/batch ranking smoke tests ---
# These give the batch endpoint a real, verifiable score spread against
# sample_jd.txt (Senior Backend Engineer: required Python/FastAPI/Django/
# PostgreSQL/Redis/Docker/Kubernetes/AWS/CI-CD; preferred Terraform/ML/
# Kafka/GraphQL).

STRONG_MATCH_LINES = [
    "Alice Chen",
    "alice.chen@example.com | (650) 555-1234 | linkedin.com/in/alicechen | github.com/alicechen",
    "",
    "Summary",
    "Senior backend engineer with 8+ years of experience building distributed systems.",
    "",
    "Experience",
    "Staff Backend Engineer, Globex Corp",
    "Mar 2016 - Present",
    "- Built and scaled REST APIs in Python using FastAPI and Django, deployed on AWS with Docker and Kubernetes.",
    "- Owned PostgreSQL schema design and Redis caching for services handling 50k req/s.",
    "- Automated infrastructure with Terraform and CI/CD pipelines using GitHub Actions.",
    "- Streamed event data through Apache Kafka and exposed internal APIs via GraphQL.",
    "- Applied Machine Learning models for fraud detection in the payments pipeline.",
    "",
    "Education",
    "Master of Science in Computer Science, Stanford University, 2015",
    "",
    "Skills",
    "Python, FastAPI, Django, PostgreSQL, Redis, Docker, Kubernetes, AWS, Terraform,",
    "GitHub Actions, CI/CD, Kafka, GraphQL, Machine Learning, Git",
]

PARTIAL_MATCH_LINES = [
    "Bob Martinez",
    "bob.martinez@example.com | (312) 555-9876",
    "",
    "Summary",
    "Backend developer with 3 years of experience.",
    "",
    "Experience",
    "Backend Developer, Initech",
    "Jun 2021 - Present",
    "- Built internal tools in Python and deployed them with Docker on AWS EC2.",
    "- Wrote REST APIs and integrated with a MySQL database.",
    "",
    "Education",
    "Bachelor of Science in Information Technology, Ohio State University, 2021",
    "",
    "Skills",
    "Python, Docker, AWS, MySQL, Git, Linux",
]

MISMATCH_LINES = [
    "Carol Nguyen",
    "carol.nguyen@example.com | (206) 555-4321",
    "",
    "Summary",
    "Junior frontend designer with 1 year of experience.",
    "",
    "Experience",
    "Junior UI/UX Designer, Creative Studio",
    "Aug 2024 - Present",
    "- Designed responsive layouts in Figma and built prototypes with Adobe XD.",
    "- Implemented designs in HTML, CSS, and JavaScript with React.",
    "",
    "Education",
    "Bachelor of Fine Arts in Graphic Design, Rhode Island School of Design, 2023",
    "",
    "Skills",
    "Figma, Adobe XD, Adobe Photoshop, HTML, CSS, React, UI/UX Design",
]


def generate_strong_match_docx():
    doc = Document()
    for line in STRONG_MATCH_LINES:
        doc.add_paragraph(line)
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, "strong_match.docx")
    doc.save(path)
    print(f"Wrote {path}")


def generate_partial_match_pdf():
    doc = fitz.open()
    page = doc.new_page()
    text = "\n".join(PARTIAL_MATCH_LINES)
    page.insert_textbox(fitz.Rect(36, 36, 560, 780), text, fontsize=10)
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, "partial_match.pdf")
    doc.save(path)
    doc.close()
    print(f"Wrote {path}")


def generate_mismatch_pdf():
    doc = fitz.open()
    page = doc.new_page()
    text = "\n".join(MISMATCH_LINES)
    page.insert_textbox(fitz.Rect(36, 36, 560, 780), text, fontsize=10)
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, "mismatch.pdf")
    doc.save(path)
    doc.close()
    print(f"Wrote {path}")


if __name__ == "__main__":
    generate_docx()
    generate_pdf()
    generate_thin_pdf()
    generate_strong_match_docx()
    generate_partial_match_pdf()
    generate_mismatch_pdf()
